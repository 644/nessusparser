from plugins import selects
from docx import Document
from docx.shared import Inches
import os.path

def genr(cb, plugin_ids, description):
	pluginid_strings = [x for x in plugin_ids if not isinstance(x, int)]
	pluginid_ints = [x for x in plugin_ids if isinstance(x, int)]

	all_rows = []

	if not len(pluginid_strings) == 0:
		all_rows += selects.pluginName(pluginid_strings, cb)
	if not len(pluginid_ints) == 0:
		all_rows += selects.pluginId(pluginid_ints, cb)
	
	all_rows = [*{*all_rows}]

	if len(all_rows) > 0:
		if os.path.isfile(cb + '.docx'):
			document = Document(cb + '.docx')
		else:
			document = Document()

		document.add_heading('Issue Description:', level=1)
		document.add_paragraph(description)
		document.save(cb + '.docx')
