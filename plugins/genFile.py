from plugins import selects
from docx import Document
from docx.shared import Inches
import os.path
from tqdm import *

def gen_document(cb, name, description, risk_description, recommendation, notes, affected_hosts):
	if os.path.isfile(cb + '.docx'):
		document = Document(cb + '.docx')
	else:
		document = Document()

	if description:
		table = document.add_table(rows=1, cols=2)
		tbl_cells = table.rows[0].cells
		tbl_cells[0].text = name
		tbl_cells = table.add_row().cells
		tbl_cells[0].text = 'Risk Rating:'
		document.add_heading('Issue Description:', level=1)
		document.add_paragraph(description)

	if recommendation:
		document.add_heading('Recommendation:', level=1)
		document.add_paragraph(recommendation)

	table = document.add_table(rows=1, cols=1)
	tbl_cells = table.rows[0].cells
	tbl_cells[0].text = 'Affected Hosts:'

	for ac in tqdm(affected_hosts):
		tbl_cells = table.add_row().cells
		#if '<bold_italic>' in ac:
		#	ac = ac.replace('<bold_italic>', '')
		#	ac = ac.replace('</bold_italic>', '')
		#	tbl_cells[0].text = str(ac) # TODO: make bold and italic
		#else:
		tbl_cells[0].text = ac
	print('\n')

	document.add_page_break()
	document.save(cb + '.docx')

def genr(cb, plugin_ids, name, description, risk_description, recommendation, notes):
	pluginid_strings = [x for x in plugin_ids if not isinstance(x, int)]
	pluginid_ints = [x for x in plugin_ids if isinstance(x, int)]

	all_rows = []
	if len(pluginid_strings) != 0:
		all_rows += selects.pluginName(pluginid_strings, cb)
	if len(pluginid_ints) != 0:
		all_rows += selects.pluginId(pluginid_ints, cb)
	
	all_rows = [*{*all_rows}]
	affected_hosts = []
	for x in all_rows:
		if x[0]:
			affected_hosts.append(x[0])
		else:
			affected_hosts.append(x[1])
	
	if len(affected_hosts) > 0:
		gen_document(cb, name, description, risk_description, recommendation, notes, affected_hosts)
